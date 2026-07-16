import streamlit as st
import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import re
import time
import random

# --- FUNGSI ANIMASI LOTTIE ---
def load_lottieurl(url: str):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()

# --- KONFIGURASI SISTEM ---
st.set_page_config(
    page_title="MI MIFTAHUSSALAM - KBC Premium",
    page_icon="❤️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- PREMIUM ADAPTIVE GLASSMORPHISM CSS ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700&display=swap');
    
    html, body, [class*="css"] { 
        font-family: 'Plus Jakarta Sans', sans-serif; 
    }
    
    .premium-card {
        background: rgba(255, 255, 255, 0.04);
        border-radius: 16px;
        padding: 28px;
        border: 1px solid rgba(255, 255, 255, 0.08);
        box-shadow: 0 10px 30px rgba(0, 0, 0, 0.05);
        backdrop-filter: blur(12px);
        -webkit-backdrop-filter: blur(12px);
        margin-bottom: 25px;
    }
    
    /* 1. Mengubah Tombol Eksekusi di Bawah Konten */
    div.stButton > button:first-child {
        background: linear-gradient(135deg, #FF4B4B 0%, #FF7676 100%);
        color: white;
        border: none;
        padding: 18px 36px;
        border-radius: 12px;
        font-size: 1.2rem;
        font-weight: 700;
        letter-spacing: 1px;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        box-shadow: 0 6px 20px rgba(255, 75, 75, 0.3);
        width: 100%;
    }
    
    div.stButton > button:first-child:hover {
        transform: translateY(-3px);
        box-shadow: 0 8px 25px rgba(255, 75, 75, 0.5);
    }

    /* 2. FIX UTAMA: MENGUBAH TAB MENJADI SHAPES TOMBOL KUSTOM YANG BESAR & TEBAL */
    div[data-testid="stTabs"] button {
        background-color: rgba(255, 255, 255, 0.05) !important; 
        border: 1px solid rgba(255, 255, 255, 0.1) !important;  
        border-radius: 30px !important;                          
        padding: 12px 28px !important;                          
        margin-right: 12px !important;                          
        transition: all 0.3s ease !important;
    }
    
    div[data-testid="stTabs"] button:hover {
        background-color: rgba(255, 75, 75, 0.1) !important;
        border-color: #FF7676 !important;
    }
    
    div[data-testid="stTabs"] button[aria-selected="true"] {
        background: linear-gradient(135deg, #FF4B4B 0%, #FF7676 100%) !important; 
        color: white !important;
        border: none !important;
        box-shadow: 0 4px 15px rgba(255, 75, 75, 0.3) !important;
    }
    
    div[data-testid="stTabs"] button [data-testid="stMarkdownContainer"] p {
        font-size: 1.25rem !important;  
        font-weight: 700 !important;    
    }
    </style>
    """, unsafe_allow_html=True)

# Konstanta API (MENGGUNAKAN SECRETS STREAMLIT CLOUD)
NVIDIA_API_KEY = st.secrets["NVIDIA_API_KEY"]
NVIDIA_API_URL = "https://integrate.api.nvidia.com/v1/chat/completions"
MODEL_NAME = "nvidia/nemotron-3-ultra-550b-a55b"

# --- LOGIKA CORE & HELPER ---
def get_ai_response_kbc(prompt, system_instruction):
    headers = {"Authorization": f"Bearer {NVIDIA_API_KEY}", "Content-Type": "application/json"}
    payload = {
        "model": MODEL_NAME,
        "messages": [
            {"role": "system", "content": system_instruction},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.3, # REVISI: DIKURANGI AGAR TIDAK HALUSINASI FORMAT
        "top_p": 0.9, 
        "max_tokens": 16384, 
        "stream": False
    }
    try:
        with st.status(f"❤️ Lagos AI Sedang Memasak... tunggu 15-45s...", expanded=True) as status:
            response = requests.post(NVIDIA_API_URL, headers=headers, json=payload, timeout=300)
            if response.status_code == 200:
                status.update(label="✅ Perumusan Konten Selesai!", state="complete", expanded=False)
                return response.json()['choices'][0]['message']['content']
            else:
                st.error(f"API Error: {response.status_code}")
                status.update(label="❌ Gagal", state="error")
                return None
    except Exception as e:
        st.error(f"Koneksi Bermasalah: {str(e)}")
        return None

def set_font_safe(run, font_name='Times New Roman', size=12, bold=False):
    if run is None: return run
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    return run

def clean_markdown_symbols(text):
    if not text: return ""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'^#+\s*', '', text)
    text = re.sub(r'^[-*]\s+', '', text)
    
    # REVISI: Mengubah <br> menjadi Enter (Baris Baru) dengan aman tanpa memecah tabel
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    
    return text.strip()

def set_document_to_two_columns(section):
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')  
    cols.set(qn('w:space'), '500')  
    sectPr.append(cols)

def set_section_to_landscape(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

# ==========================================
# FORMAT 1: MODUL AJAR / RPP / KBC MASTER (1 KOLOM + TTD KEPSEK)
# ==========================================
def create_word_doc_kbc(content, doc_type, school_data):
    try:
        doc = Document()
        
        # Atur halaman ke Landscape jika tipe dokumennya Prota, Promes, atau Jurnal
        if doc_type in ["Prota", "Promes", "Jurnal Mengajar"]:
            for section in doc.sections:
                set_section_to_landscape(section)
                section.top_margin = Inches(0.7)
                section.bottom_margin = Inches(0.7)
                section.left_margin = Inches(0.7)
                section.right_margin = Inches(0.7)

        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        style.paragraph_format.space_after = Pt(6)

        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_safe(p1.add_run(f"PEMERINTAH {school_data['kabupaten'].upper()}\n"), size=14, bold=True)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_safe(p2.add_run(f"KEMENTRIAN AGAMA ISLAM\nMADRASAH {school_data['jenis'].upper()} {school_data['nama_madrasah'].upper()}\n"), size=12, bold=True)

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_safe(p3.add_run(f"Alamat: {school_data['alamat']} | Telp: {school_data['telp']}"), size=10)

        p_line = doc.add_paragraph()
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_safe(p_line.add_run("_" * (100 if doc_type in ["Prota", "Promes", "Jurnal Mengajar"] else 70)), size=10)

        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run_title1 = p_title.add_run(f"{doc_type.upper()}\n")
        set_font_safe(run_title1, size=14, bold=True)
        run_title1.underline = True
        
        run_title2 = p_title.add_run(f"Materi: {school_data['mapel']} - {school_data['kelas']}")
        set_font_safe(run_title2, size=14, bold=True)
        run_title2.underline = True

        doc.add_paragraph()

        lines = content.split('\n')
        table_buffer = []
        in_table_mode = False
        for line in lines:
            stripped = line.strip()
            if stripped.startswith('|') and not re.match(r'^\|[\s\-:]+\|$', stripped):
                in_table_mode = True
            if in_table_mode:
                if re.match(r'^\|[\s\-:]+\|$', stripped): continue
                if stripped.startswith('|') and stripped.endswith('|'):
                    table_buffer.append([clean_markdown_symbols(c.strip()) for c in stripped.split('|')[1:-1]])
                    continue
                else:
                    if table_buffer:
                        try:
                            max_cols = max(len(row) for row in table_buffer)
                            table = doc.add_table(rows=len(table_buffer), cols=max_cols)
                            table.style = 'Table Grid'
                            for r_idx, row_data in enumerate(table_buffer):
                                row = table.rows[r_idx]
                                for c_idx, cell_text in enumerate(row_data):
                                    if c_idx < len(row.cells):
                                        row.cells[c_idx].text = cell_text
                                        for p in row.cells[c_idx].paragraphs:
                                            for run in p.runs: set_font_safe(run, size=10 if doc_type in ["Prota", "Promes", "Jurnal Mengajar"] else 11, bold=(r_idx==0))
                        except: pass
                        table_buffer = []
                        in_table_mode = False
            if not in_table_mode:
                if not stripped:
                    doc.add_paragraph()
                    continue
                p = doc.add_paragraph()
                is_bold, font_size = False, 12
                if stripped.startswith('# '):
                    stripped, is_bold, font_size = stripped[2:], True, 14
                elif stripped.startswith('## '):
                    stripped, is_bold, font_size = stripped[3:], True, 13
                if stripped.startswith('- ') or stripped.startswith('* '):
                    p.style = 'List Bullet'
                    stripped = stripped[2:]
                set_font_safe(p.add_run(clean_markdown_symbols(stripped)), size=font_size, bold=is_bold)

        doc.add_paragraph("\n")
        sig_table = doc.add_table(rows=1, cols=2)
        p_kiri = sig_table.cell(0, 0).paragraphs[0]
        p_kiri.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_safe(p_kiri.add_run(f"Mengetahui,\nKepala Madrasah\n\n\n\n\n"), size=12)
        set_font_safe(p_kiri.add_run(f"( {school_data['kepala_madrasah']} )\nNIP. {school_data['nip_kepala']}"), size=12, bold=True)

        p_kanan = sig_table.cell(0, 1).paragraphs[0]
        p_kanan.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_safe(p_kanan.add_run(f"{school_data['kota']}, {school_data['tanggal_buat']}\nGuru Mata Pelajaran\n\n\n\n\n"), size=12)
        set_font_safe(p_kanan.add_run(f"( {school_data['nama_guru']} )\nNIP. {school_data['nip_guru']}"), size=12, bold=True)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"Error Pembuatan Dokumen Ajar: {e}")
        return None

# ==========================================
# FORMAT 2: KHUSUS SOAL UJIAN REVISI (DUAL KOLOM + TANPA TTD)
# ==========================================
def create_word_soal_kbc(content, doc_type, school_data):
    try:
        doc = Document()
        current_section = doc.sections[0]
        current_section.top_margin = Inches(0.5)
        current_section.bottom_margin = Inches(0.5)
        current_section.left_margin = Inches(0.5)
        current_section.right_margin = Inches(0.5)

        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(2)
        
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_safe(p_title.add_run(f"MADRASAH IBTIDAIYAH {school_data['nama_madrasah'].replace('MI ', '').upper()}\n"), size=14, bold=True)
        set_font_safe(p_title.add_run(f"{doc_type.upper()}\n"), size=14, bold=True)
        set_font_safe(p_title.add_run(f"TAHUN PELAJARAN {school_data['tahun_ajaran']}"), size=14, bold=True)
        
        set_font_safe(doc.add_paragraph().add_run("_" * 88), size=10, bold=True)
        p_l2 = doc.add_paragraph()
        p_l2.paragraph_format.space_after = Pt(8)
        set_font_safe(p_l2.add_run("_" * 88), size=10, bold=True)
        
        table_id = doc.add_table(rows=3, cols=2)
        labels_kiri = [
            f"Nama            : ....................................",
            f"Hari / Tanggal  : ....................................",
            f"Kelas           : {school_data['kelas']}"
        ]
        labels_kanan = [
            f"No Absen      : ....................................",
            f"Waktu         : 60 Menit",
            f"Mata Pelajaran: {school_data['mapel']}"
        ]
        for idx in range(3):
            row = table_id.rows[idx]
            set_font_safe(row.cells[0].paragraphs[0].add_run(labels_kiri[idx]), size=11)
            set_font_safe(row.cells[1].paragraphs[0].add_run(labels_kanan[idx]), size=11)

        doc.add_paragraph()

        soal_section = doc.add_section()
        set_document_to_two_columns(soal_section)

        lines = content.split('\n')
        for line in lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue
            p = doc.add_paragraph()
            is_bold, font_size = False, 10.5
            if stripped.startswith('# '):
                stripped, is_bold, font_size = stripped[2:], True, 11.5
            elif stripped.startswith('## '):
                stripped, is_bold, font_size = stripped[3:], True, 10.5
            if stripped.startswith('- ') or stripped.startswith('* '):
                p.style = 'List Bullet'
                stripped = stripped[2:]
            set_font_safe(p.add_run(clean_markdown_symbols(stripped)), size=font_size, bold=is_bold)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"Error Pembuatan Soal Word: {e}")
        return None

# --- DIALOG POP-UP DOKUMEN ---
@st.dialog("📥 Berkas Anda Telah Siap!")
def show_download_popup(buffer, filename, fallback_msg):
    st.write(f"### 🎉 Berkas Selesai!")
    
    kbc_quotes = [
        "Kita semua baik-baik saja, sampai hidup menarik orang lain untuk mulai ambil suara.",
        "Mengajari hati untuk bersyukur adalah tugas seumur hidup yang tidak akan pernah selesai kita pelajari.",
        "Rasa sedih benci objek yang terus bergerak.",
        "Hati manusia itu terbatas, pilih emosi mana yang layak kamu jaga.",
        "Dunia bakal tetap bergerak, tanpa peduli berapa banyak jam tidur yang kamu punya.",
        "Semoga kita selalu punya banyak cara untuk bahagia, walau sederhana.",
        "Berterima kasih atas hal-hal yang menumbuh tumbangkan, sebab dari hal-hal itu ada makna yang mampu mendewasakan.",
        "Meskipun, kamu kalah dalam banyak hal, mari tetap menjadi manusia yang tidak habis DAYA JUANGNYA.",
        "Hiduplah dengan cara yang paling bahagia yang kamu punya dan bisa.",
    ]
    st.markdown(f"> *\"{random.choice(kbc_quotes)}\"*")
    
    st.divider()
    st.download_button(
        label="📥 UNDUH SEKARANG",
        data=buffer,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

# --- UI INTERFACE CORE ---
lottie_heart = load_lottieurl("https://assets5.lottiefiles.com/packages/lf20_v9t630.json")

c_title, c_logo = st.columns([0.85, 0.15])
with c_title:
    st.title("MI MIFTAHUSSALAM Generated Document ❤️")
    st.caption("Aplikasi Administrasi Lembaga & Bank Soal Otomatis Based Lagos AI 9.1")
with c_logo:
    if lottie_heart: st_lottie(lottie_heart, height=80, key="main_heart")

st.divider()

tab_setup, tab_materi, tab_soal = st.tabs([
    "🏛️ Langkah 1: Profil Lembaga", 
    "📚 Langkah 2: Dokumen Administrasi", 
    "📝 Langkah 3: Modul Bank Soal"
])

with tab_setup:
    st.markdown('<div class="premium-card">', unsafe_allow_html=True)
    st.subheader("Informasi Instansi & Guru")
    col1, col2 = st.columns(2)
    with col1:
        nama_madrasah = st.text_input("Nama Sekolah/Madrasah", "MI MIFTAHUSSALAM")
        jenis = st.selectbox("Jenjang Satuan", ["MI/SD", "MTs/SMP", "MA/SMA"])
        kabupaten = st.text_input("Kabupaten/Kota Domisili", "Kota Bogor")
        alamat = st.text_area("Alamat Lengkap", "Jl. Rimba Mulya II No.46, Pasirmulya, Bogor Barat", height=68)
        telp = st.text_input("Telp", "")
    with col2:
        nama_guru = st.text_input("Nama Lengkap Guru (Penyusun)", "Guru Pengajar")
        nip_guru = st.text_input("NIP Guru", "-")
        kepala_madrasah = st.text_input("Nama Kepala Madrasah", "Drs. Andi Supriadi")
        nip_kepala = st.text_input("NIP Kepala", "-")
        kota = st.text_input("Kota Pembuatan Dokumen", "Bogor")
    st.markdown('</div>', unsafe_allow_html=True)

with tab_materi:
    st.markdown('<div class="premium-card">', unsafe_allow_html=True)
    st.subheader("Kurikulum & Agenda Mengajar")
    cx, cy = st.columns(2)
    with cx:
        # MENU BARU DITAMBAHKAN DI SINI: JURNAL MENGAJAR
        doc_type_materi = st.selectbox("Pilih Output Dokumen Perangkat", [
            "Modul Ajar", "RPP", "ATP", "CP", "Prota", "Promes", "LKPD", "Jurnal Mengajar"
        ])
        mapel_materi = st.text_input("Mata Pelajaran", "Akidah Akhlak", key="mp_materi")
    with cy:
        kelas_materi = st.text_input("Kelas / Fase Tingkat", "IV / Fase B", key="kls_materi")
        tahun_ajaran_materi = st.text_input("Periode Tahun Ajaran", "2026/2027", key="ta_materi")
    materi_pembelajaran = st.text_area("Topik Pokok Pembelajaran / Bab Bahasan", "Sifat Wajib Allah", height=100, key="mat_materi")
    st.divider()
    btn_materi = st.button("🚀 GENERASIKAN DOKUMEN", type="primary", use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)

with tab_soal:
    st.markdown('<div class="premium-card">', unsafe_allow_html=True)
    st.subheader("Konfigurasi Evaluasi Ujian / Ulangan Harian (Dual Kolom)")
    
    c_soal1, c_soal2 = st.columns(2)
    with c_soal1:
        doc_type_soal = st.selectbox(
            "Jenis Asesmen Ulangan", 
            ["Sumatif Akhir Semester (SAS) Ganjil", "Sumatif Akhir Semester (SAS) Genap", "Asesmen Tengah Semester (ATS)", "Sumatif Akhir Tahun (SAT)"]
        )
        mapel_soal = st.text_input("Mata Pelajaran Ujian", "Bahasa Arab", key="mp_soal")
        kelas_soal = st.text_input("Jenjang / Kelas / Fase", "1 (Satu)", key="kls_soal")
    with c_soal2:
        materi_soal = st.text_input("Materi Pokok Ujian", "Perkenalan / Ta'aruf", key="mat_soal")
        kesulitan_soal = st.selectbox("Tingkat Kesulitan", ["Mudah", "Sedang", "Sulit", "Campuran / HOTS"])
        tahun_ajaran_soal = st.text_input("Tahun Ajaran Ujian", "2026/2027", key="ta_soal")
        
    st.markdown("##### 📊 Komposisi Distribusi Soal")
    cg1, cg2, cg3 = st.columns(3)
    with cg1:
        jml_pg = st.number_input("Jumlah Pilihan Ganda (PG)", min_value=0, max_value=50, value=10)
    with cg2:
        jml_isian = st.number_input("Jumlah Soal Isian Singkat", min_value=0, max_value=30, value=5)
    with cg3:
        jml_uraian = st.number_input("Jumlah Soal Uraian / Essay", min_value=0, max_value=20, value=2)
        
    st.divider()
    btn_soal = st.button("📝 GENERASIKAN LEMBAR BANK SOAL", type="primary", use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)

# --- SIDEBAR KONTROL ADMINISTRASI ---
with st.sidebar:
    st.header("📋 Kontrol Administrasi")
    tanggal_buat = st.date_input("Tanggal Cetak Dokumen")
    
    st.divider()
    st.write("Kurikulum Berbasis Cinta menekankan 5 Pilar Utama:")
    st.markdown("""
    - 💖 Cinta kepada Allah & Rasul
    - 📚 Cinta kepada Ilmu
    - 🧘 Cinta kepada Diri Sendiri
    - 🤝 Cinta kepada Sesama
    - 🌿 Cinta kepada Lingkungan
    """)

# --- INISIALISASI SESSION STATE ---
if "buffer_materi" not in st.session_state:
    st.session_state.buffer_materi = None
if "filename_materi" not in st.session_state:
    st.session_state.filename_materi = ""
    
if "buffer_soal" not in st.session_state:
    st.session_state.buffer_soal = None
if "filename_soal" not in st.session_state:
    st.session_state.filename_soal = ""

if "buka_popup_materi" not in st.session_state:
    st.session_state.buka_popup_materi = False
if "buka_popup_soal" not in st.session_state:
    st.session_state.buka_popup_soal = False

# --- PROSES EKSEKUSI: TAB MATERI ---
if btn_materi:
    if not nama_madrasah or not materi_pembelajaran:
        st.warning("⚠️ Harap lengkapi Profil Instansi pada Langkah 1 dan Materi pada Langkah 2.")
    else:
        data = {
            "nama_madrasah": nama_madrasah, "jenis": jenis, "kabupaten": kabupaten, "alamat": alamat, "telp": telp,
            "kepala_madrasah": kepala_madrasah, "nip_kepala": nip_kepala, "nama_guru": nama_guru, "nip_guru": nip_guru,
            "kota": kota, "tanggal_buat": tanggal_buat.strftime("%d %B %Y"),
            "mapel": mapel_materi, "kelas": kelas_materi, "tahun_ajaran": tahun_ajaran_materi, "materi": materi_pembelajaran
        }
        
        # LOGIKA PERBEDAAN FORMAT SECARA KETAT (REVISI FEW-SHOT TEMPLATE TERMASUK JURNAL)
        if doc_type_materi in ["Modul Ajar", "RPP"]:
            spesifikasi_format = (
                "Format Wajib RPP/Modul Ajar:\n"
                "Sertakan komponen formal: A. Identitas, B. Kesiapan Siswa, C. Tema Kurikulum Berbasis Cinta 5 Pilar, D. Karakteristik, E. Profil Pelajar Pancasila.\n"
                "WAJIB gunakan format tabel Markdown ini untuk langkah kegiatan:\n"
                "| Tahap Pembelajaran | Rincian Kegiatan (Pendahuluan/Inti/Penutup) | Alokasi Waktu | Nilai KBC yang Ditanamkan |\n"
                "|---|---|---|---|\n"
                "| Pendahuluan | ... | ... | ... |\n"
                "| Kegiatan Inti | ... | ... | ... |\n"
                "| Asesmen/Penilaian | ... | | ... |\n"
                "| Penutup | ... | ... | ... |\n"
            )
        elif doc_type_materi == "CP":
            spesifikasi_format = (
                "Format Wajib Capaian Pembelajaran (CP):\n"
                "Jangan buat tabel kegiatan mengajar harian! Fokus pada rasionalisasi dan target kompetensi.\n"
                "WAJIB gunakan format tabel Markdown ini untuk elemen CP:\n"
                "| Elemen CP | Deskripsi Capaian Pembelajaran | Integrasi 5 Pilar KBC |\n"
                "|---|---|---|\n"
                "| ... | ... | ... |\n"
            )
        elif doc_type_materi == "ATP":
            spesifikasi_format = (
                "Format Wajib Alur Tujuan Pembelajaran (ATP):\n"
                "WAJIB gunakan format tabel Markdown ini secara persis:\n"
                "| No | Elemen CP | Tujuan Pembelajaran (TP) | Alokasi Waktu / JP | Aspek Karakter KBC |\n"
                "|---|---|---|---|---|\n"
                "| 1 | ... | ... | ... | ... |\n"
            )
        elif doc_type_materi == "Prota":
            spesifikasi_format = (
                "Format Wajib Program Tahunan (Prota) - ORIENTASI LANDSCAPE:\n"
                "WAJIB gunakan format tabel Markdown ini secara persis:\n"
                "| No | Semester | Capaian Pembelajaran (CP) & Elemen | Tujuan Pembelajaran (TP) | Alokasi JP |\n"
                "|---|---|---|---|---|\n"
                "| 1 | Ganjil | ... | ... | ... |\n"
            )
        elif doc_type_materi == "Promes":
            spesifikasi_format = (
                "Format Wajib Program Semester (Promes) - ORIENTASI LANDSCAPE:\n"
                "WAJIB gunakan format tabel Markdown ini secara persis (Gunakan 'v' untuk menandai bulan):\n"
                "| No | Elemen CP | ATP | JP | Jul | Agu | Sep | Okt | Nov | Des | Jan | Feb | Mar | Apr | Mei | Jun |\n"
                "|---|---|---|---|---|---|---|---|---|---|---|---|---|---|---|---|\n"
                "| 1 | ... | ... | ... | v | v |   |   |   |   |   |   |   |   |   |   |\n"
            )
        elif doc_type_materi == "Jurnal Mengajar":
            spesifikasi_format = (
                "Format Wajib Jurnal Agenda Mengajar Harian Guru - ORIENTASI LANDSCAPE:\n"
                "Jabarkan kegiatan untuk beberapa kali pertemuan sesuai topik.\n"
                "WAJIB gunakan format tabel Markdown ini secara persis untuk log harian:\n"
                "| Pertemuan Ke | Hari/Tanggal | Kelas/Jam Ke | Kompetensi Dasar / TP | Materi Pokok | Uraian Kegiatan | Absensi/Keterangan |\n"
                "|---|---|---|---|---|---|---|\n"
                "| 1 | ... | ... | ... | ... | ... | ... |\n"
                "| 2 | ... | ... | ... | ... | ... | ... |\n"
            )
        else:  # LKPD
            spesifikasi_format = (
                "Format Wajib LKPD KBC:\n"
                "Buat lembar kerja aktivitas interaktif (Judul, Petunjuk, Pertanyaan Pemantik, Studi Kasus).\n"
                "WAJIB gunakan format tabel Markdown ini untuk rubrik penilaian di bagian akhir:\n"
                "| Aspek Penilaian | Kriteria Ketercapaian | Skor Maksimal |\n"
                "|---|---|---|\n"
                "| ... | ... | ... |\n"
            )

        sys_prompt = (
            f"Anda adalah Pakar AI dan Ahli Kurikulum Merdeka terintegrasi Deep learning (meaningful,mindful,joyful).\n"
            f"Tugas utama Anda menyusun dokumen: {doc_type_materi}.\n"
            f"PERATURAN UTAMA: Anda harus mematuhi struktur berikut dan DILARANG keluar format:\n"
            f"PERATURAN Kedua: Anda harus Memasukan minimal satu Mindful,Meaningful,joyful di kegiatan inti pembelajaran dan sesuaikan dengan materi:\n"
            f"{spesifikasi_format}\n"
            f"Gunakan format Markdown murni. Jangan gunakan tag HTML seperti <br> atau <p>. Gunakan format standar (\\n)."
            f"Gunakan CP berdasarkan KMA Nomor 503 Tah1un 2025.\n"
            f"Masukan Nilai KKM semua mata pelajaran 78.\n"
        )
        
        user_prompt = f"""
        Buat dokumen perangkat: **{doc_type_materi}**
        
        DATA IDENTITAS (WAJIB DIMASUKKAN KE DALAM DOKUMEN):
        - Nama Instansi / Sekolah: {nama_madrasah}
        - Nama Guru Penyusun: {nama_guru}
        - Kepala Madrasah: {kepala_madrasah}
        
        DETAIL PELAJARAN:
        - Mata Pelajaran: {mapel_materi}
        - Kelas/Fase: {kelas_materi}
        - Tahun Ajaran: {tahun_ajaran_materi}
        - Topik/Bab Utama: {materi_pembelajaran}
        
        Pastikan format Anda persis seperti yang diminta di instruksi sistem! 
        Sertakan Data Identitas di atas secara lengkap pada bagian "Identitas Modul" atau "Informasi Umum" di hasil akhir.
        """
        
        content = get_ai_response_kbc(user_prompt, sys_prompt)
        if content:
            buffer = create_word_doc_kbc(content, doc_type_materi, data)
            if buffer:
                st.session_state.buffer_materi = buffer
                st.session_state.filename_materi = f"KBC2026_{doc_type_materi.replace(' ', '_')}_{mapel_materi}.docx"
                st.session_state.buka_popup_materi = True
                st.rerun()

# --- PROSES EKSEKUSI: TAB SOAL ---
if btn_soal:
    if not nama_madrasah or not materi_soal:
        st.warning("⚠️ Harap lengkapi instansi pada Langkah 1 dan detail komposisi soal pada Langkah 3.")
    else:
        data = {
            "nama_madrasah": nama_madrasah, "jenis": jenis, "kabupaten": kabupaten, "alamat": alamat, 
            "kepala_madrasah": kepala_madrasah, "nip_kepala": nip_kepala, 
            "nama_guru": nama_guru, "nip_guru": nip_guru, 
            "kota": kota, "tanggal_buat": tanggal_buat.strftime("%d %B %Y"),
            "mapel": mapel_soal, "kelas": kelas_soal, "tahun_ajaran": tahun_ajaran_soal, 
            "materi": materi_soal, "tingkat_kesulitan": kesulitan_soal
        }
        
        sys_prompt = """
        Anda adalah Ahli Evaluasi Akademik Kurikulum Merdeka Deep Learning berbasis Cinta (KBC 2026)
        Tugas Anda memproduksi naskah soal ujian. Pastikan menyentuh nilai-nilai karakter 5 pilar cinta KBC.
        Format Output Wajib:
        - BAGIAN I: SOAL PILIHAN GANDA (opsi A, B, C, D)
        - BAGIAN II: SOAL ISIAN SINGKAT
        - BAGIAN III: SOAL URAIAN / ESSAY
        Buatkan Kunci Jawaban terpisah di paling bawah dokumen. Jangan gunakan blok kode (```).
        """
        
        user_prompt = f"""
        Ujian {doc_type_soal} Mapel {mapel_soal}, Kelas {kelas_soal}.
        Materi: {materi_soal}. Tingkat Kesulitan: {kesulitan_soal}.
        Komposisi:
        - {jml_pg} Pilihan Ganda (PG)
        - {jml_isian} Isian Singkat
        - {jml_uraian} Uraian/Essay
        Sertakan Kunci Jawaban lengkap di akhir!
        """
        
        content = get_ai_response_kbc(user_prompt, sys_prompt)
        if content:
            buffer = create_word_soal_kbc(content, doc_type_soal, data)
            if buffer:
                st.session_state.buffer_soal = buffer
                st.session_state.filename_soal = f"Soal_{doc_type_soal.split(' ')[0]}_{mapel_soal}_{kelas_soal.replace(' ', '')}.docx"
                st.session_state.buka_popup_soal = True
                st.rerun()

# --- LOGIKA PEMANGGILAN POPUP ---
if st.session_state.buka_popup_materi:
    st.session_state.buka_popup_materi = False  
    show_download_popup(st.session_state.buffer_materi, st.session_state.filename_materi, "Pekerjaan selesai.")

if st.session_state.buka_popup_soal:
    st.session_state.buka_popup_soal = False  
    show_download_popup(st.session_state.buffer_soal, st.session_state.filename_soal, "Pekerjaan selesai.")
